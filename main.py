from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import StreamingResponse
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import docx2txt
import os
import json
import asyncio
from dotenv import load_dotenv
from groq import Groq, AsyncGroq
from pydantic import BaseModel

load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError("API_KEY_ERROR")

# Sync client: used for resume parsing (a one-off, not streamed).
client = Groq(api_key=my_api_key)
# Async client: used for the chat answer, so it can be streamed to the
# browser as real, incremental chunks instead of arriving all at once.
async_client = AsyncGroq(api_key=my_api_key)
model = "openai/gpt-oss-120b"

# Accepted resume formats, in the order they're searched for.
SUPPORTED_RESUME_EXTENSIONS = (".pdf", ".docx")

# Optional: set this to an exact filename (e.g. "My_Resume.docx") if you
# want to force a specific file. Leave it empty to auto-detect — the app
# will use the first .pdf or .docx file it finds next to main.py.
RESUME_FILE = os.getenv(
    "RESUME_FILE",
    "./private/Mukesh_Kumar_Resume.docx"
)

# Accepted behavioral-document formats, in the order they're searched for.
# This document holds STAR-format stories / behavioral Q&A prep material
# (e.g. "Tell me about a time you faced conflict at work") so the chatbot
# can answer HR behavioral questions, not just resume/skills questions.
SUPPORTED_BEHAVIORAL_EXTENSIONS = (".pdf", ".docx", ".txt")

# Optional: set this to an exact filename (e.g. "Behavioral_Answers.docx")
# to force a specific file. Leave it empty to auto-detect — the app will
# use the first file next to main.py whose name contains "behav"
# (case-insensitive), matching .pdf, .docx, or .txt.
BEHAVIORAL_FILE = os.getenv(
    "BEHAVIORAL_FILE",
    "./private/HR_Behavioural_Interview_Answer_Bank_Chatbot_document.docx"
)

app = FastAPI(title="Candidate Chatbot API")

# Allow the frontend (any origin) to call this API. Needed any time the
# frontend is opened from a different origin than the backend, e.g. during
# local development or if you host frontend/backend on different domains.
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skill_used: list[str] = []


class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experience: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []


resume_schema = Resume.model_json_schema()


class ChatRequest(BaseModel):
    question: str


# ---------------------------------------------------------------------------
# The resume PDF is parsed by the LLM ONCE and cached in memory, instead of
# being re-read and re-parsed on every single chat message. This is what
# keeps each HR question fast (one LLM call instead of two) and keeps your
# Groq usage low.
# ---------------------------------------------------------------------------
_resume_cache: Resume | None = None

# The behavioral document is plain prep text (STAR stories, behavioral Q&A,
# etc.) — unlike the resume, it's NOT parsed into structured JSON, since
# behavioral answers are narrative and free-form. It's just read once and
# cached as raw text, then handed to the model as extra context.
_behavioral_cache: str | None = None
_behavioral_loaded: bool = False


def read_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(file_path: Path) -> str:
    """
    Extract text from a .docx file.

    python-docx's `document.paragraphs` only walks body-level paragraphs —
    it does NOT see text placed inside text boxes, which many designed
    resume templates use for the candidate's name/header. docx2txt parses
    the underlying XML more broadly and does pick up text-box content, so
    we run it first and fall back to (or merge in) python-docx output for
    anything it might miss (e.g. some table layouts).
    """
    text = ""
    try:
        text = docx2txt.process(str(file_path)) or ""
    except Exception:
        text = ""

    document = Document(file_path)
    fallback_lines = []
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            fallback_lines.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    fallback_lines.append(cell.text)
    fallback_text = "\n".join(fallback_lines)

    # Merge both extractions so nothing is lost: docx2txt tends to catch
    # text boxes/headers that python-docx misses, while python-docx is
    # more reliable for some table structures.
    combined = (text + "\n" + fallback_text).strip()
    return combined if combined else fallback_text


def read_txt(file_path: Path) -> str:
    return file_path.read_text(encoding="utf-8", errors="ignore")


def read_resume_text(file_path: Path) -> str:
    """Extract text from a resume file, dispatching by extension."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(file_path)
    if suffix == ".docx":
        return read_docx(file_path)
    raise HTTPException(
        status_code=500,
        detail=f"Unsupported resume file type '{suffix}'. Use .pdf or .docx.",
    )


def read_document_text(file_path: Path) -> str:
    """Extract text from a document, dispatching by extension (pdf/docx/txt)."""
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return read_pdf(file_path)
    if suffix == ".docx":
        return read_docx(file_path)
    if suffix == ".txt":
        return read_txt(file_path)
    raise HTTPException(
        status_code=500,
        detail=f"Unsupported file type '{suffix}'. Use .pdf, .docx, or .txt.",
    )


def guess_name_from_text(text: str) -> str | None:
    """
    Deterministic fallback for the candidate's name: on the vast majority
    of resumes, the very first non-empty line IS the name (before any
    contact info or section heading). Used only when the LLM parse comes
    back with name=null, so a flaky/odd model response doesn't lose
    something that's plainly sitting at the top of the document.
    """
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        # Reject lines that are clearly not a bare name: emails, phone
        # numbers/dates, or long section text.
        if "@" in line or any(ch.isdigit() for ch in line):
            return None
        words = line.split()
        if 1 <= len(words) <= 5 and len(line) <= 60:
            return line
        return None
    return None


def find_resume_file() -> Path:
    """
    Resolve which resume file to use:
    1. If RESUME_FILE is set, use it exactly (error if missing).
    2. Otherwise, auto-detect next to main.py — preferring a filename that
       contains "resume" or "cv", and always skipping anything that looks
       like the behavioral document (contains "behav"), so the two auto
       detectors can never pick up each other's file.
    """
    if RESUME_FILE:
        path = Path(RESUME_FILE)
        if not path.exists():
            raise HTTPException(
                status_code=500,
                detail=f"Resume file '{RESUME_FILE}' was not found next to main.py.",
            )
        return path

    here = Path(".")
    candidates = []
    for ext in SUPPORTED_RESUME_EXTENSIONS:
        for match in sorted(here.glob(f"*{ext}")):
            if "behav" in match.stem.lower():
                continue
            candidates.append(match)

    if not candidates:
        raise HTTPException(
            status_code=500,
            detail="No resume file found. Add a .pdf or .docx resume next to main.py "
            "(or set RESUME_FILE to an exact filename).",
        )

    # Prefer a filename that explicitly says "resume" or "cv" over an
    # arbitrary other document that happens to sort first alphabetically.
    for match in candidates:
        stem = match.stem.lower()
        if "resume" in stem or stem.endswith("cv") or "_cv_" in stem or "-cv-" in stem:
            return match

    return candidates[0]


def find_behavioral_file() -> Path | None:
    """
    Resolve which behavioral document to use:
    1. If BEHAVIORAL_FILE is set, use it exactly (error if missing).
    2. Otherwise, auto-detect a file next to main.py whose name contains
       "behav" (case-insensitive) with a supported extension.
    3. If nothing is found, return None — the behavioral doc is optional,
       so the app still runs fine without one (it just won't have
       behavioral-specific prep material to draw on).
    """
    if BEHAVIORAL_FILE:
        path = Path(BEHAVIORAL_FILE)
        if not path.exists():
            raise HTTPException(
                status_code=500,
                detail=f"Behavioral file '{BEHAVIORAL_FILE}' was not found next to main.py.",
            )
        return path

    here = Path(".")
    for ext in SUPPORTED_BEHAVIORAL_EXTENSIONS:
        for match in sorted(here.glob(f"*{ext}")):
            if "behav" in match.stem.lower():
                return match

    return None


def parse_resume(resume_text: str) -> Resume:
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    6. The candidate's name is usually the most prominent line near the
       very top of the document (often larger/bold text, a header, or a
       text box) — it may appear before any section heading like
       "Experience" or "Contact". Look there first for the "name" field.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt},
    ]
    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"},
    )
    data = json.loads(response.choices[0].message.content)
    return Resume(**data)


def get_resume() -> Resume:
    """Return the cached, parsed resume, parsing it on first use only."""
    global _resume_cache
    if _resume_cache is None:
        path = find_resume_file()
        text = read_resume_text(path)
        if not text.strip():
            raise HTTPException(
                status_code=500,
                detail=f"Could not extract any text from '{path.name}'. "
                "It may be a scanned/image-based document.",
            )
        _resume_cache = parse_resume(text)
        # Safety net: if the LLM didn't extract a name even though it's
        # plainly the first line of the document, fill it in directly
        # instead of showing a blank/placeholder name in the UI.
        if not _resume_cache.name or not _resume_cache.name.strip():
            guessed = guess_name_from_text(text)
            if guessed:
                _resume_cache.name = guessed
    return _resume_cache


def get_behavioral_content() -> str | None:
    """
    Return the cached behavioral-document text, reading it on first use
    only. Returns None if no behavioral document was found — this is
    optional, so its absence is not an error.
    """
    global _behavioral_cache, _behavioral_loaded
    if not _behavioral_loaded:
        _behavioral_loaded = True
        path = find_behavioral_file()
        if path is not None:
            text = read_document_text(path)
            if text.strip():
                _behavioral_cache = text
    return _behavioral_cache


async def ask_candidate_stream(question: str, resume: Resume, behavioral_content: str | None):
    """
    Stream the candidate answer while keeping the Groq request safely below
    the current token-per-minute limit.

    The behavioral document can be much larger than the structured resume.
    It is therefore included only for questions that are actually behavioral,
    and it is capped to a reasonable size. This prevents the full behavioral
    document from being sent with every question.
    """

    # Behavioral material is only useful for questions that ask for a
    # situation/story/HR example. Do not send the whole behavioral document
    # for normal skills, experience, project, or "tell me about yourself"
    # questions.
    behavioral_keywords = (
        "tell me about a time",
        "describe a time",
        "give me an example",
        "example of a",
        "conflict",
        "challenge",
        "difficult situation",
        "failure",
        "mistake",
        "disagreement",
        "leadership",
        "teamwork",
        "team conflict",
        "pressure",
        "stress",
        "deadline",
        "stakeholder",
        "feedback",
        "weakness",
        "strength",
        "behavioral",
        "situation",
        "problem you faced",
        "how did you handle",
        "how do you handle",
    )

    question_lower = question.lower()
    is_behavioral = any(keyword in question_lower for keyword in behavioral_keywords)

    # Keep a substantial safety margin below Groq's current 8,000 TPM limit.
    # 12,000 characters is roughly 3,000 tokens for typical English text.
    MAX_BEHAVIORAL_CHARS = 12000

    if behavioral_content and is_behavioral:
        behavioral_context = behavioral_content[:MAX_BEHAVIORAL_CHARS]

        behavioral_section = f"""
Below is the candidate's behavioral interview prep material (STAR-format
stories, past behavioral Q&A, etc.). Use it to answer the behavioral question.

Behavioral prep:
{behavioral_context}
"""
    else:
        behavioral_section = """
No behavioral prep material is included for this question. Use the structured
resume information below for normal questions about experience, skills,
projects, education, and career background.
"""

    system_prompt = f"""
You are an AI assistant representing a job candidate.

Below is the structured information available about the candidate's resume:

{resume.model_dump_json(indent=2)}

{behavioral_section}

Rules:
1. Answer only using the information provided above.
2. Never hallucinate or invent experience, skills, employers, dates, or projects.
3. If information is unavailable, say:
   "I don't have enough information to answer that."
4. Be professional and concise.
5. Answer as if HR is interviewing the candidate.
6. For behavioral questions, prefer the STAR format (Situation, Task, Action,
   Result) when the behavioral prep material supports it.
"""

    print(
        f"Chat request | behavioral={is_behavioral} | "
        f"behavioral_chars={len(behavioral_context) if behavioral_content and is_behavioral else 0}"
    )

    try:
        stream = await async_client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": question},
            ],
            stream=True,
        )

        async for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta.content
            if delta:
                yield delta

    except Exception as e:
        # Log the real error so it is visible in the backend/Render logs.
        # Do not hide Groq errors behind the misleading "not enough information"
        # message.
        print(f"LLM ERROR: {type(e).__name__}: {e}")
        yield "I’m temporarily unable to answer because the AI service returned an error. Please try again."


@app.get("/api/profile")
def profile():
    """A small public summary the frontend uses to render the profile card."""
    resume = get_resume()
    return {
        "name": resume.name,
        "total_experience_years": resume.total_experience_years,
        "skills": resume.skills[:10],
        "education": resume.education,
        "has_behavioral_doc": get_behavioral_content() is not None,
    }


# ---------------------------------------------------------------------------
# TEMPORARY DEBUG ENDPOINT — remove once the name/profile issue is confirmed
# fixed. Lets you see exactly what text was extracted from the resume file
# and what the LLM parsed from it, so you can tell whether the problem is
# extraction (text missing) or parsing (text present but not captured).
# ---------------------------------------------------------------------------
@app.get("/api/debug/resume")
def debug_resume():
    path = find_resume_file()
    raw_text = read_resume_text(path)
    resume = get_resume()
    return {
        "resume_file": path.name,
        "extracted_text_length": len(raw_text),
        "extracted_text_preview": raw_text[:1500],
        "parsed_resume": resume.model_dump(),
    }


@app.post("/chat")
async def chat(request: ChatRequest):
    # Resolve/parse the resume BEFORE streaming starts, so a parsing error
    # comes back as a normal error response instead of breaking mid-stream.
    # get_resume() is blocking (sync), so it's run off the event loop —
    # this only matters on the very first request; every request after
    # that hits the in-memory cache and returns instantly.
    resume = await asyncio.to_thread(get_resume)
    behavioral_content = await asyncio.to_thread(get_behavioral_content)
    return StreamingResponse(
        ask_candidate_stream(request.question, resume, behavioral_content),
        media_type="text/plain",
        headers={
            # Ask any reverse proxy in front of this (Render, nginx, etc.)
            # not to buffer the response — without this, some proxies hold
            # the whole reply and send it in one go, killing the streaming
            # effect even though the server is sending it incrementally.
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",
        },
    )


# Serve the frontend (frontend/index.html) at "/".
# IMPORTANT: this must be the LAST route registered — routes defined above
# are matched first, and anything not matched falls through to these files.
app.mount("/", StaticFiles(directory="frontend", html=True), name="frontend")